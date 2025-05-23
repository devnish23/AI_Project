# services/document_processor.py - Document Processing Service
import os
import hashlib
import aiofiles
from typing import List, Dict, Any
from pathlib import Path
import asyncio
import logging

# Document parsing libraries
import PyPDF2
import docx
from pdfplumber import PDF
import tiktoken

# Text processing
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles document upload, parsing, chunking, and embedding generation.
    
    Key Features:
    - Supports PDF, DOCX, TXT, MD files
    - Semantic text chunking with overlap
    - Metadata extraction and preservation
    - Async file processing for performance
    """
    
    def __init__(self):
        self.upload_dir = Path("data/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize text splitter for semantic chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,        # Target chunk size in characters
            chunk_overlap=200,      # Overlap between chunks for context
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]  # Priority separators
        )
        
        # Initialize embedding model
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Token counter for cost estimation
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    async def save_uploaded_file(self, file) -> str:
        """
        Save uploaded file to local storage.
        
        Args:
            file: UploadFile object from FastAPI
            
        Returns:
            str: Path to saved file
        """
        try:
            # Generate unique filename to avoid conflicts
            file_hash = hashlib.md5(file.filename.encode()).hexdigest()[:8]
            safe_filename = f"{file_hash}_{file.filename}"
            file_path = self.upload_dir / safe_filename
            
            # Save file asynchronously
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            logger.info(f"File saved: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Failed to save file {file.filename}: {str(e)}")
            raise
    
    async def process_document(self, file_path: str, document_type: str, user_id: str) -> Dict[str, Any]:
        """
        Process document: extract text, create chunks, generate embeddings.
        
        Args:
            file_path: Path to the document file
            document_type: Type of document (policy, sop, log, etc.)
            user_id: ID of the user who uploaded the document
            
        Returns:
            Dict containing processing results
        """
        try:
            # Extract text based on file type
            extracted_text = await self._extract_text(file_path)
            
            # Create semantic chunks
            chunks = self._create_chunks(extracted_text, file_path)
            
            # Generate embeddings for chunks
            embeddings = self._generate_embeddings(chunks)
            
            # Store in vector database
            from services.rag_service import RAGService
            rag_service = RAGService()
            
            chunk_ids = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = rag_service.store_chunk(
                    content=chunk['content'],
                    embedding=embedding,
                    metadata={
                        **chunk['metadata'],
                        'document_type': document_type,
                        'user_id': user_id,
                        'chunk_index': i
                    }
                )
                chunk_ids.append(chunk_id)
            
            # Calculate some statistics
            token_count = len(self.encoding.encode(extracted_text))
            
            result = {
                'chunks_created': len(chunks),
                'total_tokens': token_count,
                'chunk_ids': chunk_ids,
                'extracted_text_length': len(extracted_text)
            }
            
            logger.info(f"Document processed: {file_path}, {len(chunks)} chunks created")
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_path}: {str(e)}")
            raise
    
    async def _extract_text(self, file_path: str) -> str:
        """
        Extract text from various file formats.
        
        Args:
            file_path: Path to the document
            
        Returns:
            str: Extracted text content
        """
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                return await self._extract_docx_text(file_path)
            elif file_extension in ['.txt', '.md']:
                return await self._extract_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods for better accuracy."""
        text_content = ""
        
        try:
            # Method 1: Try pdfplumber (better for complex layouts)
            with open(file_path, 'rb') as file:
                with PDF(file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            # Method 2: Fallback to PyPDF2 if pdfplumber fails
            if not text_content.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text or markdown files."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
                return content
                
        except Exception as e:
            logger.error(f"Text file extraction failed: {str(e)}")
            raise
    
    def _create_chunks(self, text: str, file_path: str) -> List[Dict[str, Any]]:
        """
        Create semantic chunks from extracted text.
        
        Args:
            text: Extracted text content
            file_path: Original file path for metadata
            
        Returns:
            List of chunk dictionaries with content and metadata
        """
        try:
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text)
            
            # Create chunk objects with metadata
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = {
                    'content': chunk_text.strip(),
                    'metadata': {
                        'source_file': os.path.basename(file_path),
                        'full_path': file_path,
                        'chunk_index': i,
                        'total_chunks': len(text_chunks),
                        'char_count': len(chunk_text),
                        'token_count': len(self.encoding.encode(chunk_text))
                    }
                }
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Chunk creation failed: {str(e)}")
            raise
    
    def _generate_embeddings(self, chunks: List[Dict[str, Any]]) -> List[List[float]]:
        """
        Generate embeddings for text chunks using sentence transformer.
        
        Args:
            chunks: List of chunk dictionaries
            
        Returns:
            List of embedding vectors
        """
        try:
            # Extract text content from chunks
            texts = [chunk['content'] for chunk in chunks]
            
            # Generate embeddings in batch for efficiency
            embeddings = self.embedding_model.encode(
                texts,
                show_progress_bar=True,
                convert_to_tensor=False,
                normalize_embeddings=True  # L2 normalization for cosine similarity
            )
            
            # Convert to list of lists for JSON serialization
            return embeddings.tolist()
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {str(e)}")
            raise
    
    async def delete_document(self, document_id: str):
        """
        Delete document file and clean up associated data.
        
        Args:
            document_id: Unique identifier for the document
        """
        try:
            # Get document metadata from database
            from models.database import Database
            db = Database()
            document = db.get_document_by_id(document_id)
            
            if document and document['file_path']:
                # Delete physical file
                file_path = Path(document['file_path'])
                if file_path.exists():
                    file_path.unlink()
                    logger.info(f"Deleted file: {file_path}")
            
        except Exception as e:
            logger.error(f"Document deletion failed: {str(e)}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Return list of supported file formats."""
        return ['.pdf', '.docx', '.txt', '.md']
    
    def estimate_processing_time(self, file_size_mb: float) -> float:
        """
        Estimate processing time based on file size.
        
        Args:
            file_size_mb: File size in megabytes
            
        Returns:
            Estimated processing time in seconds
        """
        # Rough estimates based on file type and size
        # PDF: ~30 seconds per MB, DOCX: ~15 seconds per MB, TXT: ~5 seconds per MB
        return max(5.0, file_size_mb * 20)  # Minimum 5 seconds
    
    async def batch_process_documents(self, file_paths: List[str], document_type: str, user_id: str):
        """
        Process multiple documents concurrently.
        
        Args:
            file_paths: List of file paths to process
            document_type: Type classification for all documents
            user_id: User who owns the documents
            
        Returns:
            List of processing results
        """
        try:
            # Create semaphore to limit concurrent processing
            semaphore = asyncio.Semaphore(3)  # Process max 3 files concurrently
            
            async def process_single(file_path):
                async with semaphore:
                    return await self.process_document(file_path, document_type, user_id)
            
            # Process all files concurrently
            tasks = [process_single(file_path) for file_path in file_paths]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Separate successful results from exceptions
            successful = []
            failed = []
            
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    failed.append({
                        'file_path': file_paths[i],
                        'error': str(result)
                    })
                else:
                    successful.append({
                        'file_path': file_paths[i],
                        'result': result
                    })
            
            logger.info(f"Batch processing complete: {len(successful)} successful, {len(failed)} failed")
            
            return {
                'successful': successful,
                'failed': failed,
                'total_processed': len(successful)
            }
            
        except Exception as e:
            logger.error(f"Batch processing failed: {str(e)}")
            raise